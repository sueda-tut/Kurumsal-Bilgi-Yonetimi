# PDF, Word ve Excel dosyalarından temiz metin çıkarır

from dataclasses import dataclass
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from openpyxl import load_workbook


DESTEKLENEN_UZANTILAR = {".pdf", ".docx", ".xlsx"}


@dataclass(frozen=True, slots=True)
class MetinBolumu:
    metin: str
    sayfa_no: int | None = None
    sayfa_adi: str | None = None
    satir_no: int | None = None


class MetinCikarmaHatasi(Exception):
    """Dosyadan metin çıkarılamadığında oluşan hata."""


def metni_temizle(metin: str) -> str:
    """Gereksiz boşlukları temizleyerek metni standartlaştırır."""

    satirlar = []

    for satir in metin.splitlines():
        temiz_satir = " ".join(satir.split())

        if temiz_satir:
            satirlar.append(temiz_satir)

    return "\n".join(satirlar)


def hucre_degerini_metne_cevir(deger: Any) -> str:
    """Excel hücre değerini temiz bir metne dönüştürür."""

    if deger is None:
        return ""

    if isinstance(deger, float) and deger.is_integer():
        return str(int(deger))

    return metni_temizle(str(deger))


def pdf_metni_cikar(dosya_yolu: str | Path) -> list[MetinBolumu]:
    """PDF dosyasını sayfa numaraları korunacak şekilde okur."""

    dosya_yolu = Path(dosya_yolu)
    bolumler: list[MetinBolumu] = []

    try:
        with fitz.open(dosya_yolu) as pdf:
            for sayfa_indeksi, sayfa in enumerate(pdf):
                metin = metni_temizle(sayfa.get_text("text"))

                if metin:
                    bolumler.append(
                        MetinBolumu(
                            metin=metin,
                            sayfa_no=sayfa_indeksi + 1,
                        )
                    )

    except Exception as hata:
        raise MetinCikarmaHatasi(
            f"PDF dosyasından metin çıkarılamadı: {dosya_yolu.name}"
        ) from hata

    return bolumler


def word_metni_cikar(dosya_yolu: str | Path) -> list[MetinBolumu]:
    """Word belgesindeki paragrafları ve tabloları metne dönüştürür."""

    dosya_yolu = Path(dosya_yolu)
    bolumler: list[MetinBolumu] = []

    try:
        belge = Document(dosya_yolu)

        for paragraf in belge.paragraphs:
            metin = metni_temizle(paragraf.text)

            if metin:
                bolumler.append(MetinBolumu(metin=metin))

        for tablo in belge.tables:
            for satir_no, satir in enumerate(tablo.rows, start=1):
                hucreler = [
                    metni_temizle(hucre.text)
                    for hucre in satir.cells
                    if metni_temizle(hucre.text)
                ]

                if hucreler:
                    bolumler.append(
                        MetinBolumu(
                            metin=" | ".join(hucreler),
                            satir_no=satir_no,
                        )
                    )

    except Exception as hata:
        raise MetinCikarmaHatasi(
            f"Word dosyasından metin çıkarılamadı: {dosya_yolu.name}"
        ) from hata

    return bolumler


def excel_metni_cikar(dosya_yolu: str | Path) -> list[MetinBolumu]:
    """Excel satırlarını 'Sütun: değer' biçiminde metne dönüştürür."""

    dosya_yolu = Path(dosya_yolu)
    bolumler: list[MetinBolumu] = []

    try:
        calisma_kitabi = load_workbook(
            filename=dosya_yolu,
            read_only=True,
            data_only=True,
        )

        for sayfa in calisma_kitabi.worksheets:
            satirlar = sayfa.iter_rows(values_only=True)
            baslik_satiri = next(satirlar, None)

            if baslik_satiri is None:
                continue

            basliklar = [
                hucre_degerini_metne_cevir(deger)
                or f"Sütun {indeks}"
                for indeks, deger in enumerate(
                    baslik_satiri,
                    start=1,
                )
            ]

            for satir_no, satir in enumerate(satirlar, start=2):
                alanlar = []

                for baslik, deger in zip(basliklar, satir):
                    metin_degeri = hucre_degerini_metne_cevir(deger)

                    if metin_degeri:
                        alanlar.append(
                            f"{baslik}: {metin_degeri}"
                        )

                if alanlar:
                    bolumler.append(
                        MetinBolumu(
                            metin=" | ".join(alanlar),
                            sayfa_adi=sayfa.title,
                            satir_no=satir_no,
                        )
                    )

        calisma_kitabi.close()

    except Exception as hata:
        raise MetinCikarmaHatasi(
            f"Excel dosyasından metin çıkarılamadı: {dosya_yolu.name}"
        ) from hata

    return bolumler


def dosyadan_metin_cikar(
    dosya_yolu: str | Path,
) -> list[MetinBolumu]:
    """Dosya uzantısına göre uygun metin çıkarıcıyı çalıştırır."""

    dosya_yolu = Path(dosya_yolu)

    if not dosya_yolu.exists():
        raise FileNotFoundError(
            f"Dosya bulunamadı: {dosya_yolu}"
        )

    uzanti = dosya_yolu.suffix.lower()

    if uzanti == ".pdf":
        return pdf_metni_cikar(dosya_yolu)

    if uzanti == ".docx":
        return word_metni_cikar(dosya_yolu)

    if uzanti == ".xlsx":
        return excel_metni_cikar(dosya_yolu)

    raise ValueError(
        "Desteklenmeyen dosya türü. "
        "Desteklenen türler: pdf, docx, xlsx."
    )