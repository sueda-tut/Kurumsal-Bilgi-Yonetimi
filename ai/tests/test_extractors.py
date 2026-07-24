# PDF, Word ve Excel metin çıkarıcılarını örnek dosyalarla test eder

import fitz
from docx import Document
from openpyxl import Workbook

from ai.extractors import dosyadan_metin_cikar


def test_pdf_metni_sayfa_numaralariyla_cikarilir(tmp_path):
    pdf_yolu = tmp_path / "ornek.pdf"

    pdf = fitz.open()

    birinci_sayfa = pdf.new_page()
    birinci_sayfa.insert_text(
        (72, 72),
        "Birinci sayfa kurumsal bilgi.",
    )

    ikinci_sayfa = pdf.new_page()
    ikinci_sayfa.insert_text(
        (72, 72),
        "Ikinci sayfa personel bilgisi.",
    )

    pdf.save(pdf_yolu)
    pdf.close()

    bolumler = dosyadan_metin_cikar(pdf_yolu)

    assert len(bolumler) == 2
    assert bolumler[0].sayfa_no == 1
    assert bolumler[1].sayfa_no == 2
    assert "Birinci sayfa" in bolumler[0].metin
    assert "Ikinci sayfa" in bolumler[1].metin


def test_word_metni_cikarilir(tmp_path):
    word_yolu = tmp_path / "ornek.docx"

    belge = Document()
    belge.add_heading("Personel El Kitabı", level=1)
    belge.add_paragraph(
        "Çalışanlar şirket kurallarına uymalıdır."
    )
    belge.save(word_yolu)

    bolumler = dosyadan_metin_cikar(word_yolu)
    birlesik_metin = "\n".join(
        bolum.metin for bolum in bolumler
    )

    assert "Personel El Kitabı" in birlesik_metin
    assert "Çalışanlar şirket kurallarına uymalıdır." in (
        birlesik_metin
    )


def test_excel_satirlari_sutun_degerine_donusturulur(
    tmp_path,
):
    excel_yolu = tmp_path / "ornek.xlsx"

    calisma_kitabi = Workbook()
    sayfa = calisma_kitabi.active
    sayfa.title = "Personeller"

    sayfa.append(["Ad Soyad", "Departman", "Aktif"])
    sayfa.append(["Ahmet Yılmaz", "İnsan Kaynakları", True])
    sayfa.append(["Ayşe Demir", "Muhasebe", False])

    calisma_kitabi.save(excel_yolu)
    calisma_kitabi.close()

    bolumler = dosyadan_metin_cikar(excel_yolu)

    assert len(bolumler) == 2
    assert bolumler[0].sayfa_adi == "Personeller"
    assert bolumler[0].satir_no == 2
    assert "Ad Soyad: Ahmet Yılmaz" in bolumler[0].metin
    assert "Departman: İnsan Kaynakları" in bolumler[0].metin
    assert "Aktif: True" in bolumler[0].metin